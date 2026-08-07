import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from pathlib import Path

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 54, 93) # Deep Navy
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(41, 128, 185) # Blue
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)
    return p

def add_body_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def add_bullet_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def format_table_headers(table, col_widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, cell in enumerate(hdr_cells):
        set_cell_background(cell, "1B365D")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(10.5)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
    
    for row_idx, row in enumerate(table.rows[1:], start=1):
        bg_color = "F4F6F7" if row_idx % 2 == 1 else "FFFFFF"
        for i, cell in enumerate(row.cells):
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(10)

def build_full_srs():
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ═══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    r = p.add_run("SIGNBRIDGE: A DUAL-COMMUNICATION ROBOTIC TRANSLATOR FOR INDIAN SIGN LANGUAGE (ISL)")
    r.font.name = 'Calibri'
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(27, 54, 93)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("A Minor Project Report - Software Requirement Specification")
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Submitted By:")
    r.font.name = 'Calibri'
    r.font.size = Pt(12)
    r.font.bold = True

    students = [
        ("(23BT04076)", "Stuti Mistry"),
        ("(24BT04D225)", "Foram Patel"),
        ("(24BT04D231)", "Chintan Sharma"),
        ("(24BT04D236)", "Parth Thakkar"),
    ]
    for roll, name in students:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{roll:<18}")
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        r2 = p.add_run(name)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(11)
        r2.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("In partial fulfilment of the requirement for the award of the degree of\nBachelor of Technology In\nSemester VI - Computer Science and Engineering (AI-ML)")
    r.font.name = 'Calibri'
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Department of Computer Science and Engineering\nSchool of Technology\nGSFC University,\nVigyan Bhavan, P. O. Fertilizer Nagar,\nVadodara - 391750, Gujarat, India\nMay 2026")
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.bold = True

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CERTIFICATE & DECLARATION
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "Certificate")
    add_body_p(doc, "This is to certify that the report submitted along with the project entitled 'SIGNBRIDGE: A DUAL-COMMUNICATION ROBOTIC TRANSLATOR FOR INDIAN SIGN LANGUAGE (ISL)' has been carried out by Stuti Mistry (23BT04076), Foram Patel (24BT04D225), Chintan Sharma (24BT04D231), and Parth Thakkar (24BT04D236) under supervision for the partial fulfillment of the degree of Bachelor of Technology in Computer Science and Engineering (AI-ML) at GSFC University, Vadodara.")
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.add_run("SIGNATURE (Guide / HOD)\t\t\t\tSIGNATURE (Internal Examiner)\nDate: May, 2026\t\t\t\t\tDate: May, 2026")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # LIST OF TABLES & FIGURES
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "List Of Tables")
    t0 = doc.add_table(rows=1, cols=2)
    t0.rows[0].cells[0].paragraphs[0].text = "Table No."
    t0.rows[0].cells[1].paragraphs[0].text = "Title"
    tables_list = [
        ("Table 1.1", "Hardware and Software Specifications"),
        ("Table 2.1", "Comparative Analysis of Existing Technical Solutions and Their Drawbacks"),
        ("Table 2.2", "Literature Review Summary of Key Research Papers"),
        ("Table 3.1", "Abbreviations and Acronyms Reference"),
        ("Table 5.1", "Software Testing Matrix and Verification Objectives"),
        ("Table 6.1", "Future Improvements and Extension Roadmap"),
    ]
    for row in tables_list:
        r = t0.add_row()
        r.cells[0].paragraphs[0].text = row[0]
        r.cells[1].paragraphs[0].text = row[1]
    format_table_headers(t0)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    add_heading_1(doc, "List Of Figures")
    t1 = doc.add_table(rows=1, cols=2)
    t1.rows[0].cells[0].paragraphs[0].text = "Figure No."
    t1.rows[0].cells[1].paragraphs[0].text = "Title"
    figures_list = [
        ("Figure 1.1", "SignBridge Dual-Communication Loop Architecture"),
        ("Figure 3.1", "System Use Case Diagram"),
        ("Figure 3.2", "System Activity Diagram (Input & Output Pipelines)"),
        ("Figure 4.1", "MediaPipe 21 Hand Landmark Topology"),
        ("Figure 4.2", "SignBridge Touchscreen Kiosk User Interface"),
    ]
    for row in figures_list:
        r = t1.add_row()
        r.cells[0].paragraphs[0].text = row[0]
        r.cells[1].paragraphs[0].text = row[1]
    format_table_headers(t1)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "TABLE OF CONTENTS")
    toc_items = [
        ("CHAPTER 1: INTRODUCTION", "1"),
        ("  1.1 Project Description", "1"),
        ("      Problem Statement", "1"),
        ("      Objectives", "2"),
        ("  1.2 Features", "2"),
        ("  1.3 Tools and Technology Used", "3"),
        ("  1.4 Requirements for Hardware and Software", "4"),
        ("CHAPTER 2: LITERATURE REVIEW", "5"),
        ("  2.1 Origin of the Problem & History of the Topic", "5"),
        ("  2.2 Comparative Study of Existing Technical Solutions and Their Drawbacks", "6"),
        ("CHAPTER 3: SYSTEM DESIGN", "8"),
        ("  3.1 Use Case", "8"),
        ("  3.2 Activity Diagram", "10"),
        ("CHAPTER 4: IMPLEMENTATION OF PROJECT", "12"),
        ("  4.1 System Architecture & Data Flow", "12"),
        ("  4.2 Module Implementation", "13"),
        ("CHAPTER 5: SOFTWARE TESTING", "16"),
        ("  5.1 Rationale for Testing", "16"),
        ("  5.2 Levels of Testing", "16"),
        ("  5.3 Testing Methods", "17"),
        ("      5.3.1 Functionality Testing", "17"),
        ("      5.3.2 UX & Feedback", "18"),
        ("CHAPTER 6: LIMITATIONS AND FUTURE ENHANCEMENT", "19"),
        ("  6.1 Existing Limitations", "19"),
        ("  6.2 Future Improvements", "20"),
        ("CHAPTER 7: CONCLUSION", "21"),
        ("CHAPTER 8: REFERENCES", "22"),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{item:<65} {page}")
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        if "CHAPTER" in item:
            r.font.bold = True

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 1: INTRODUCTION
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "CHAPTER 1: INTRODUCTION")
    
    add_heading_2(doc, "1.1 Project Description")
    add_heading_3(doc, "Problem Statement")
    add_body_p(doc, "The first obstacle faced when establishing accessibility for Deaf and Hard-of-Hearing individuals lies within the 'Communication Gap.' Society relies predominantly on spoken and written languages, while Deaf signers communicate through Indian Sign Language (ISL).")
    add_body_p(doc, "Unlike American Sign Language (ASL), which uses a mostly one-handed manual alphabet, Indian Sign Language is fundamentally a two-handed sign language—most letters (A–Z), digits (0–9), and practical vocabulary require synchronized movements from both hands. Existing sign language translation tools are often limited by three major flaws: (1) they are one-way only (translating signs to text or text to signs, but rarely both); (2) they rely on expensive, intrusive sensor gloves that restrict movement; and (3) they utilize single-hand robotic models that cannot physically perform two-handed ISL manual alphabets.")

    add_heading_3(doc, "Objectives")
    add_body_p(doc, "The goal of SignBridge is to develop an affordable, two-way (dual) robotic translation system that bridges the gap between Deaf signers and hearing individuals using webcam-based AI vision and low-cost robotic hands. The system aims to:")
    add_bullet_p(doc, "Provide a bidirectional translation loop (ISL signs -> Spoken Speech/Text for hearing users, and Spoken Speech/Text -> Physical ISL robotic hand gestures for Deaf signers).", "Two-Way (Dual) Communication: ")
    add_bullet_p(doc, "Eliminate wearable sensor gloves by using standard laptop webcams with Google MediaPipe 3D hand-landmark tracking.", "Non-Invasive Vision Pipeline: ")
    add_bullet_p(doc, "Build low-cost dual 5-finger robotic hands driven by 10 micro servos to physically perform ISL manual alphabets and gestures.", "Dual Scrap-Material Robotic Hands: ")
    add_bullet_p(doc, "Provide a touch-optimized, full-screen interactive interface suitable for deployment at hospital reception desks, public service counters, and educational institutions.", "Public Kiosk Accessibility: ")

    add_heading_2(doc, "1.2 Features")
    add_body_p(doc, "The system is built around a guided 'Dual-Communication' progression:")
    add_body_p(doc, "Steps to be taken include:")
    add_bullet_p(doc, "Initializing the live webcam vision pipeline, server health check, and Arduino hardware link.", "Commencing the Kiosk Session — ")
    add_bullet_p(doc, "Real-time 3D hand joint tracking, TensorFlow deep-learning classification, live text display, and spoken audio output via Text-to-Speech (TTS).", "ISL Gesture Input Path (Deaf -> Hearing) — ")
    add_bullet_p(doc, "Providing interactive touchscreen editing tools including Undo, Delete, Clear, Add Space, Copy, and Speak Aloud.", "Sentence Builder Assistance — ")
    add_bullet_p(doc, "Speech-to-text microphone audio transcription or direct text input entry.", "Speech & Text Input Path (Hearing -> Deaf) — ")
    add_bullet_p(doc, "Looking up pre-calibrated servo angle sets and transmitting serial commands to actuate the dual robotic hands in sequence.", "Robotic Hand Actuation — ")
    add_bullet_p(doc, "Rendering live side-by-side panel status for both the signer ('YOU') and the robotic translator ('ROBOT').", "Dual Display Stream — ")

    add_heading_2(doc, "1.3 Tools and Technology Used")
    add_body_p(doc, "For the project to be both powerful and user-friendly, the following development stack was used:")
    add_bullet_p(doc, "Frontend library and build tool for managing state and rendering the responsive touchscreen kiosk UI.", "React + Vite — ")
    add_bullet_p(doc, "Micro web framework hosting the backend REST API endpoints for translation model inference and serial communication.", "Python & Flask — ")
    add_bullet_p(doc, "Computer vision framework for real-time extraction of 42 3D hand landmark points (126 coordinate features per frame).", "Google MediaPipe — ")
    add_bullet_p(doc, "Machine learning platform used to train the Multi-Layer Perceptron (MLP) ISL gesture classifier.", "TensorFlow / Keras — ")
    add_bullet_p(doc, "Microcontroller board utilized for simultaneous 10-channel PWM servo motor control.", "Arduino Mega 2560 — ")
    add_bullet_p(doc, "9g micro servo motors providing tendon-pulling tension for finger movements.", "SG90 Micro Servos — ")
    add_bullet_p(doc, "Python library for USB serial data communication between the Flask backend and Arduino board.", "PySerial — ")

    add_heading_2(doc, "1.4 Requirements for Hardware and Software")
    add_bullet_p(doc, "Processor with standard CPU/GPU execution support (Intel i3/i5 or equivalent), 4GB+ RAM, Integrated/USB 720p Webcam, Microphone, Speakers, 1x Arduino Mega 2560, 10x SG90 Micro Servos, 5V 4A External Power Supply, dual scrap-material hand frames.", "Hardware — ")
    add_bullet_p(doc, "Developed using React 18, Vite 5, Python 3.10+, Flask 3.0, MediaPipe, TensorFlow 2.x, PySerial, Arduino IDE.", "Software Framework — ")
    add_bullet_p(doc, "Windows 10/11, macOS, or Linux.", "Operating System — ")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 2: LITERATURE REVIEW
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "CHAPTER 2: LITERATURE REVIEW")
    
    add_heading_2(doc, "2.1 Origin of the Problem & History of the Topic")
    add_body_p(doc, "Sign Language Recognition (SLR) research began in the 1990s with physical sensor gloves (such as DataGlove) equipped with flex sensors and accelerometers. While precise, sensor gloves proved unsuited for widespread adoption due to high costs, delicate wiring, and tactile discomfort for native signers.")
    add_body_p(doc, "With the rise of computer vision in the 2010s, researchers shifted toward camera-based recognition using Convolutional Neural Networks (CNNs). However, Indian Sign Language (ISL) presented unique structural challenges compared to Western sign languages like ASL or BSL. Most ISL letters require two hands overlapping or touching key points (e.g. index finger touching opposite palm). This structural requirement makes single-hand gesture tools insufficient for real-world ISL communication.")

    add_heading_2(doc, "2.2 Comparative Study of Existing Technical Solutions and Their Drawbacks")
    add_body_p(doc, "Below is a comparative analysis of existing communication solutions vs. the SignBridge dual-communication platform:")

    t4 = doc.add_table(rows=1, cols=5)
    hdr = t4.rows[0].cells
    hdr[0].paragraphs[0].text = "Solution"
    hdr[1].paragraphs[0].text = "Interactivity"
    hdr[2].paragraphs[0].text = "Bidirectionality"
    hdr[3].paragraphs[0].text = "2-Hand Support"
    hdr[4].paragraphs[0].text = "Drawbacks"

    comp_data = [
        ("Printed Charts / Books", "Low", "None", "Static Images", "Passive; no real-time translation or feedback."),
        ("Flex Sensor Gloves", "Medium", "1-Way (Sign -> Text)", "Rarely (High Cost)", "Intrusive wiring, fragile components, high maintenance cost."),
        ("ASL 1-Hand Robotic Models", "Medium", "1-Way (Text -> Sign)", "No (Single Hand)", "Incompatible with two-handed ISL alphabets."),
        ("SignBridge Kiosk System", "High", "2-Way Dual Loop", "Yes (Dual 10-Servo)", "Low-cost scrap materials require angle calibration."),
    ]
    for row in comp_data:
        r = t4.add_row()
        for i, val in enumerate(row):
            r.cells[i].paragraphs[0].text = val
    format_table_headers(t4)

    add_heading_3(doc, "Literature Review Summary Table")
    t5 = doc.add_table(rows=1, cols=4)
    hdr5 = t5.rows[0].cells
    hdr5[0].paragraphs[0].text = "Sr no."
    hdr5[1].paragraphs[0].text = "Paper / Dataset Title"
    hdr5[2].paragraphs[0].text = "Summary & Findings"
    hdr5[3].paragraphs[0].text = "Tools, Technologies & Methodologies"

    lit_data = [
        ("1", "INCLUDE: A Large Scale Dataset for Indian Sign Language Recognition (IISc Bangalore & AI4Bharat)", "Developed a comprehensive benchmark ISL dataset containing thousands of video clips across 263 ISL signs. Demonstrated that two-handed spatial keypoints significantly outperform raw RGB frames for ISL gesture classification.", "Tools: MediaPipe Holistic, Python, PyTorch\nMethodology: 3D Hand Landmark Extraction, Spatio-Temporal Keypoint Normalization."),
        ("2", "Mendeley ISL Benchmark Dataset: 2-Handed ISL Manual Alphabets & Posture Features", "Established a standardized 2-handed ISL alphabet dataset (A-Z) documenting left and right hand postures, finger extension ratios, and touch points for geometric heuristic and deep learning validation.", "Tools: Mendeley Data, MediaPipe Hands, Python\nMethodology: 42-point 3D keypoint mapping, posture feature extraction."),
        ("3", "Google MediaPipe: On-Device Real-Time Hand Tracking Framework", "Presents a real-time 21 3D hand landmark estimation pipeline running efficiently on low-power CPU/GPU devices without specialized depth sensors.", "Tools: MediaPipe Hands, TensorFlow Lite\nMethodology: BlazePalm detector + Hand Landmark regression network."),
    ]
    for row in lit_data:
        r = t5.add_row()
        for i, val in enumerate(row):
            r.cells[i].paragraphs[0].text = val
    format_table_headers(t5)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 3: SYSTEM DESIGN
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "CHAPTER 3. SYSTEM DESIGN")
    
    add_heading_2(doc, "3.1 Use Case")
    add_body_p(doc, "The SignBridge system involves three primary user roles / actors: Deaf Signer, Hearing Speaker, and System Administrator.")
    add_bullet_p(doc, "Signs in front of webcam; views real-time text output; uses live sentence builder controls (Undo, Delete, Clear, Speak).", "Deaf Signer — ")
    add_bullet_p(doc, "Speaks into microphone or types text message; watches dual robotic hands physically perform ISL signs.", "Hearing Speaker — ")
    add_bullet_p(doc, "Connects/disconnects Arduino Mega serial port; runs dataset importer scripts and model training pipelines.", "System Administrator — ")

    add_heading_2(doc, "3.2 Activity Diagram")
    add_body_p(doc, "The operational workflow consists of two parallel pipelines within the Dual-Communication Loop:")
    add_body_p(doc, "1. Input Path (Signer -> Hearing User): Webcam capture -> MediaPipe 42 keypoint extraction -> Landmark Normalization (wrist center & scale) -> Keras MLP Classifier inference -> Recognized text display -> Text-to-Speech audio engine.")
    add_body_p(doc, "2. Output Path (Hearing User -> Signer): Microphone speech-to-text / typed text input -> Character parsing -> Servo angle lookup table (ISL_SERVO_MAP) -> USB PySerial transmission -> Arduino Mega PWM driver -> 10x SG90 servo motor physical hand movement.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 4: IMPLEMENTATION OF PROJECT
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "CHAPTER 4. IMPLEMENTATION OF PROJECT")
    
    add_heading_2(doc, "4.1 System Architecture & Data Flow")
    add_body_p(doc, "SignBridge is structured into three integrated architectural tiers:")
    add_bullet_p(doc, "React + Vite single-page application rendering the touchscreen split kiosk UI (HumanPanel and RobotPanel).", "1. Presentation Layer (Frontend Kiosk UI): ")
    add_bullet_p(doc, "Python Flask REST API (app.py) providing endpoints for model prediction (/api/translate), dataset collection (/api/collect_data), and robot control (/api/robot/sign).", "2. Intelligence & API Layer (Backend Engine): ")
    add_bullet_p(doc, "Arduino Mega 2560 running a custom C++ sketch driving 10 PWM SG90 servos across dual scrap-material hand frames.", "3. Physical Hardware Layer (Robotic Servo Controller): ")

    add_heading_2(doc, "4.2 Module Implementation")
    add_body_p(doc, "Key implementation highlights include:")
    add_bullet_p(doc, "MediaPipe integration extracts 21 keypoints per hand in real-time.", "Vision Module (CameraView.jsx): ")
    add_bullet_p(doc, "Loads isl_gesture_model.h5 and normalizes 126-element landmark arrays before classification.", "Classifier Module (translator_model.py): ")
    add_bullet_p(doc, "Manages serial connection at 9600 baud rate and maps letters A-Z to 10-servo angle arrays.", "Arduino Serial Controller (arduino_serial.py): ")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 5: SOFTWARE TESTING
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "CHAPTER 5. SOFTWARE TESTING")
    
    add_heading_2(doc, "5.1 Rationale for Testing")
    add_body_p(doc, "Testing ensures real-time gesture recognition accuracy (>95%), low API latency (<100ms), and reliable physical servo actuation without mechanical binding.")

    add_heading_2(doc, "5.2 Levels of Testing")
    add_bullet_p(doc, "Tested MediaPipe landmark parsing, normalization functions, and PySerial message formatting.", "Unit Testing — ")
    add_bullet_p(doc, "Verified HTTP communication between React frontend and Flask API, and USB communication between Flask and Arduino.", "Integration Testing — ")
    add_bullet_p(doc, "Executed end-to-end live testing of both communication paths under varying room lighting.", "System Testing — ")

    add_heading_2(doc, "5.3 Testing Methods")
    add_heading_3(doc, "5.3.1 Functionality Testing")
    add_body_p(doc, "Verified classification accuracy across all 26 ISL letters and confirmed correct servo angles for each sign.")

    add_heading_3(doc, "5.3.2 UX & Feedback")
    add_body_p(doc, "Evaluated touchscreen button responsiveness (Undo, Delete, Clear, Speak) and clear visual panel hierarchy.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 6: LIMITATIONS AND FUTURE ENHANCEMENT
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "CHAPTER 6. LIMITATIONS AND FUTURE ENHANCEMENT")
    
    add_heading_2(doc, "6.1 Existing Limitations")
    add_bullet_p(doc, "Cardboard and plastic joints require periodic recalibration under string tension.", "1. Mechanical Wear of Scrap Materials: ")
    add_bullet_p(doc, "Extremely dark environments impact webcam landmark detection accuracy.", "2. Ambient Lighting Dependence: ")

    add_heading_2(doc, "6.2 Future Improvements")
    add_bullet_p(doc, "Replacing scrap frames with 3D-printed biomimetic hands and N20 gear motors.", "1. 3D-Printed Robotic Chassis: ")
    add_bullet_p(doc, "Integrating LSTM/Transformer models to recognize dynamic whole-word phrases.", "2. Sequence Classification for Whole Words: ")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 7 & 8: CONCLUSION & REFERENCES
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "CHAPTER 7. CONCLUSION")
    add_body_p(doc, "SignBridge successfully demonstrates a two-way (dual) robotic translation system for Indian Sign Language. By combining MediaPipe vision tracking, TensorFlow deep learning, and low-cost Arduino servo control, SignBridge provides an accessible solution bridging the communication gap.")

    add_heading_1(doc, "CHAPTER 8: REFERENCES")
    add_bullet_p(doc, "IISc Bangalore & AI4Bharat INCLUDE Dataset Paper (2020).", "1. ")
    add_bullet_p(doc, "Mendeley ISL Benchmark Dataset (2-Handed ISL Alphabets), Dataset 98mzk82wbb.", "2. ")
    add_bullet_p(doc, "Google MediaPipe Framework Documentation (mediapipe.dev).", "3. ")
    add_bullet_p(doc, "TensorFlow & Keras API Documentation (tensorflow.org).", "4. ")
    add_bullet_p(doc, "Arduino Mega 2560 Hardware & Servo Library Reference (arduino.cc).", "5. ")

    # Save to file
    out_path = Path("Software Requirement Specification.docx")
    doc.save(out_path)
    print(f"Successfully generated updated SRS document at: {out_path.resolve()}")

if __name__ == '__main__':
    build_full_srs()
